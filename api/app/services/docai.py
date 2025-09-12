"""Google Document AI integration for document structure extraction.

This module provides integration with Google Document AI for extracting
structured data from documents including tables, forms, and layout information.
"""

from __future__ import annotations

import os
import json
import logging
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass

try:
    from google.cloud import documentai_v1 as documentai
    from google.oauth2 import service_account
    HAS_DOCAI = True
except ImportError:
    HAS_DOCAI = False
    # PRODUCTION REQUIREMENT: Google Document AI must be installed and configured
    raise ImportError(
        "Google Document AI is required for production deployment. "
        "Install with: pip install google-cloud-documentai"
    )
    
    class service_account:
        class Credentials:
            @staticmethod
            def from_service_account_file(path):
                return None

from ..core.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentAIResult:
    """Result from Document AI processing."""
    text: str
    tables: List[Dict[str, Any]]
    forms: List[Dict[str, Any]]
    entities: List[Dict[str, Any]]
    layout: Dict[str, Any]
    confidence: float
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "text": self.text,
            "tables": self.tables,
            "forms": self.forms,
            "entities": self.entities,
            "layout": self.layout,
            "confidence": self.confidence
        }


class DocumentAIProcessor:
    """Service for processing documents with Google Document AI."""
    
    def __init__(self):
        """Initialize Document AI client."""
        self.client = None
        self.processor_id = os.getenv("DOC_AI_PROCESSOR_ID") or settings.doc_ai_processor_id
        
        # Initialize client if credentials are available and Document AI is installed
        if HAS_DOCAI:
            creds_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS") or settings.google_application_credentials
            if creds_path and os.path.exists(creds_path):
                try:
                    credentials = service_account.Credentials.from_service_account_file(creds_path)
                    self.client = documentai.DocumentProcessorServiceClient(credentials=credentials)
                    logger.info(f"Document AI client initialized with processor: {self.processor_id}")
                except Exception as e:
                    logger.error(f"Failed to initialize Document AI client: {e}")
                    self.client = None
            else:
                logger.warning("Google Application Credentials not configured for Document AI")
        else:
            logger.warning("Google Document AI library not installed, using fallback extraction")
    
    def process_document(
        self,
        document_bytes: bytes,
        mime_type: str = "application/pdf"
    ) -> DocumentAIResult:
        """Process a document to extract structured data.
        
        Args:
            document_bytes: Raw document bytes
            mime_type: MIME type of the document
            
        Returns:
            DocumentAIResult with extracted structure
        """
        if not self.client or not self.processor_id:
            logger.warning("Document AI not configured, returning basic extraction")
            return self._fallback_extraction(document_bytes, mime_type)
        
        try:
            # Parse processor path
            project, location, processor = self._parse_processor_path(self.processor_id)
            name = self.client.processor_path(project, location, processor)
            
            # Create the request
            raw_document = documentai.RawDocument(
                content=document_bytes,
                mime_type=mime_type
            )
            request = documentai.ProcessRequest(
                name=name,
                raw_document=raw_document
            )
            
            # Process the document
            logger.info(f"Processing document with Document AI (size: {len(document_bytes)} bytes)")
            result = self.client.process_document(request=request)
            
            # Extract structured data
            document = result.document
            
            # Extract tables
            tables = self._extract_tables(document)
            
            # Extract form fields
            forms = self._extract_forms(document)
            
            # Extract entities
            entities = self._extract_entities(document)
            
            # Extract layout information
            layout = self._extract_layout(document)
            
            # Calculate overall confidence
            confidence = self._calculate_confidence(document)
            
            logger.info(
                f"Document processed: {len(tables)} tables, {len(forms)} forms, "
                f"{len(entities)} entities, confidence: {confidence:.2f}"
            )
            
            return DocumentAIResult(
                text=document.text,
                tables=tables,
                forms=forms,
                entities=entities,
                layout=layout,
                confidence=confidence
            )
            
        except Exception as e:
            logger.error(f"Document AI processing failed: {e}")
            return self._fallback_extraction(document_bytes, mime_type)
    
    def _parse_processor_path(self, path: str) -> Tuple[str, str, str]:
        """Parse processor path into components.
        
        Expected format: projects/{project}/locations/{location}/processors/{processor}
        """
        parts = path.split("/")
        if len(parts) >= 6:
            project = parts[1]
            location = parts[3]
            processor = parts[5]
            return project, location, processor
        else:
            # Return defaults if path is malformed
            return "default-project", "us", "default-processor"
    
    def _extract_tables(self, document) -> List[Dict[str, Any]]:
        """Extract tables from the document."""
        tables = []
        
        if not hasattr(document, 'pages'):
            return tables
            
        for page in document.pages:
            if not hasattr(page, 'tables'):
                continue
                
            for table in page.tables:
                table_data = {
                    "rows": [],
                    "confidence": getattr(table, 'confidence', 0.0)
                }
                
                # Extract header rows
                if hasattr(table, 'header_rows'):
                    for row in table.header_rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = self._get_text_from_layout(document, cell.layout)
                            row_data.append(cell_text)
                        table_data["rows"].append({"type": "header", "cells": row_data})
                
                # Extract body rows
                if hasattr(table, 'body_rows'):
                    for row in table.body_rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = self._get_text_from_layout(document, cell.layout)
                            row_data.append(cell_text)
                        table_data["rows"].append({"type": "body", "cells": row_data})
                
                tables.append(table_data)
        
        return tables
    
    def _extract_forms(self, document) -> List[Dict[str, Any]]:
        """Extract form fields from the document."""
        forms = []
        
        if not hasattr(document, 'pages'):
            return forms
            
        for page in document.pages:
            if not hasattr(page, 'form_fields'):
                continue
                
            for field in page.form_fields:
                field_name = ""
                field_value = ""
                
                if hasattr(field, 'field_name') and field.field_name:
                    field_name = self._get_text_from_layout(document, field.field_name.layout)
                
                if hasattr(field, 'field_value') and field.field_value:
                    field_value = self._get_text_from_layout(document, field.field_value.layout)
                
                forms.append({
                    "name": field_name,
                    "value": field_value,
                    "confidence": getattr(field, 'confidence', 0.0)
                })
        
        return forms
    
    def _extract_entities(self, document) -> List[Dict[str, Any]]:
        """Extract named entities from the document."""
        entities = []
        
        if not hasattr(document, 'entities'):
            return entities
            
        for entity in document.entities:
            entity_data = {
                "type": getattr(entity, 'type_', 'unknown'),
                "text": getattr(entity, 'mention_text', ''),
                "confidence": getattr(entity, 'confidence', 0.0),
                "properties": []
            }
            
            if hasattr(entity, 'properties'):
                for prop in entity.properties:
                    entity_data["properties"].append({
                        "type": getattr(prop, 'type_', 'unknown'),
                        "text": getattr(prop, 'mention_text', ''),
                        "confidence": getattr(prop, 'confidence', 0.0)
                    })
            
            entities.append(entity_data)
        
        return entities
    
    def _extract_layout(self, document) -> Dict[str, Any]:
        """Extract layout information from the document."""
        layout = {
            "pages": []
        }
        
        if not hasattr(document, 'pages'):
            return layout
            
        for page in document.pages:
            page_info = {
                "width": 0,
                "height": 0,
                "blocks": 0,
                "paragraphs": 0,
                "lines": 0,
                "tokens": 0
            }
            
            if hasattr(page, 'dimension') and page.dimension:
                page_info["width"] = getattr(page.dimension, 'width', 0)
                page_info["height"] = getattr(page.dimension, 'height', 0)
            
            if hasattr(page, 'blocks'):
                page_info["blocks"] = len(page.blocks)
            if hasattr(page, 'paragraphs'):
                page_info["paragraphs"] = len(page.paragraphs)
            if hasattr(page, 'lines'):
                page_info["lines"] = len(page.lines)
            if hasattr(page, 'tokens'):
                page_info["tokens"] = len(page.tokens)
            
            layout["pages"].append(page_info)
        
        return layout
    
    def _get_text_from_layout(self, document, layout) -> str:
        """Extract text from a layout element."""
        if not layout or not hasattr(layout, 'text_anchor') or not layout.text_anchor:
            return ""
        
        text = ""
        if hasattr(layout.text_anchor, 'text_segments'):
            for segment in layout.text_anchor.text_segments:
                start = getattr(segment, 'start_index', 0) or 0
                end = getattr(segment, 'end_index', len(document.text)) or len(document.text)
                text += document.text[start:end]
        
        return text.strip()
    
    def _calculate_confidence(self, document) -> float:
        """Calculate overall confidence score for the document."""
        confidences = []
        
        # Collect all confidence scores
        if hasattr(document, 'pages'):
            for page in document.pages:
                if hasattr(page, 'tables'):
                    for table in page.tables:
                        if hasattr(table, 'confidence'):
                            confidences.append(table.confidence)
                
                if hasattr(page, 'form_fields'):
                    for field in page.form_fields:
                        if hasattr(field, 'confidence'):
                            confidences.append(field.confidence)
        
        if hasattr(document, 'entities'):
            for entity in document.entities:
                if hasattr(entity, 'confidence'):
                    confidences.append(entity.confidence)
        
        # Return average confidence or 0.5 if no scores
        if confidences:
            return sum(confidences) / len(confidences)
        return 0.5
    
    def _fallback_extraction(
        self,
        document_bytes: bytes,
        mime_type: str
    ) -> DocumentAIResult:
        """Fallback extraction when Document AI is not available."""
        # Basic text extraction
        text = ""
        if mime_type == "text/plain":
            text = document_bytes.decode('utf-8', errors='ignore')
        else:
            # For other formats, use proper document processing libraries
            try:
                if mime_type == "application/pdf":
                    import PyPDF2
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(document_bytes))
                    text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(document_bytes))
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    raise ValueError(f"Unsupported document type: {mime_type}")
            except ImportError as e:
                raise ImportError(f"Document processing requires additional libraries: {e}")
            except Exception as e:
                raise ValueError(f"Failed to process document: {e}")
        
        return DocumentAIResult(
            text=text,
            tables=[],
            forms=[],
            entities=[],
            layout={"pages": [{"text_length": len(text)}]},
            confidence=0.0
        )


# Global processor instance
_processor: Optional[DocumentAIProcessor] = None


def get_document_ai_processor() -> DocumentAIProcessor:
    """Get the global Document AI processor instance."""
    global _processor
    if _processor is None:
        _processor = DocumentAIProcessor()
    return _processor


def process_document_bytes(content: bytes, mime_type: str = "application/pdf") -> Optional[object]:
    """Legacy function for backward compatibility.
    
    Process a document with Google Document AI, returning the Document object.
    """
    processor = get_document_ai_processor()
    result = processor.process_document(content, mime_type)
    
    # Return a mock document object for backward compatibility
    if HAS_DOCAI:
        doc = documentai.Document()
        doc.text = result.text
        return doc
    else:
        doc = MockDocument()
        doc.text = result.text
        return doc


def extract_text_blocks(doc) -> List[str]:
    """Legacy function for backward compatibility.
    
    Extract text blocks from a document.
    """
    if not doc:
        return []
    
    out: List[str] = []
    
    # If it's a DocumentAIResult, extract text
    if isinstance(doc, DocumentAIResult):
        if doc.text:
            # Split into paragraphs
            paragraphs = doc.text.split('\n\n')
            out.extend(p.strip() for p in paragraphs if p.strip())
    # Otherwise handle as Document object
    elif hasattr(doc, 'pages'):
        for page in doc.pages:
            if hasattr(page, 'paragraphs'):
                for para in page.paragraphs:
                    text = _layout_to_text(para.layout, doc)
                    if text:
                        out.append(text)
    
    # Fallback to full text if no paragraphs found
    if not out and hasattr(doc, 'text') and doc.text:
        out.append(doc.text)
    
    return out


def _layout_to_text(layout, doc) -> str:
    """Legacy helper for extracting text from layout."""
    if not layout or not hasattr(layout, 'text_anchor') or not layout.text_anchor:
        return ""
    
    txt = []
    if hasattr(layout.text_anchor, 'text_segments'):
        for seg in layout.text_anchor.text_segments:
            start = int(getattr(seg, 'start_index', 0) or 0)
            end = int(getattr(seg, 'end_index', 0) or 0)
            if end > start:
                txt.append(doc.text[start:end])
    
    return "".join(txt).strip()


def process_document(
    document_bytes: bytes,
    mime_type: str = "application/pdf"
) -> Dict[str, Any]:
    """Process a document and return structured data.
    
    Args:
        document_bytes: Raw document bytes
        mime_type: MIME type of the document
        
    Returns:
        Dictionary with extracted structure
    """
    processor = get_document_ai_processor()
    result = processor.process_document(document_bytes, mime_type)
    return result.to_dict()


def extract_brand_canon_from_document(
    document_bytes: bytes,
    mime_type: str = "application/pdf"
) -> Dict[str, Any]:
    """Extract brand canon information from a document.
    
    Args:
        document_bytes: Raw document bytes
        mime_type: MIME type of the document
        
    Returns:
        Dictionary with brand canon elements
    """
    # Process the document
    doc_data = process_document(document_bytes, mime_type)
    
    # Extract brand-relevant information
    brand_canon = {
        "colors": [],
        "fonts": [],
        "logos": [],
        "guidelines": []
    }
    
    # Look for color information in forms and entities
    for form in doc_data.get("forms", []):
        name = form.get("name", "").lower()
        value = form.get("value", "")
        
        if any(color_term in name for color_term in ["color", "colour", "hex", "rgb", "pantone"]):
            brand_canon["colors"].append({
                "name": form.get("name"),
                "value": value,
                "confidence": form.get("confidence", 0.0)
            })
        
        if any(font_term in name for font_term in ["font", "typeface", "typography"]):
            brand_canon["fonts"].append({
                "name": form.get("name"),
                "value": value,
                "confidence": form.get("confidence", 0.0)
            })
    
    # Look for guidelines in the text
    text = doc_data.get("text", "")
    if "brand" in text.lower() or "guideline" in text.lower():
        # Extract sentences containing brand guidelines
        sentences = text.split(".")
        for sentence in sentences:
            if any(term in sentence.lower() for term in ["must", "should", "always", "never", "required"]):
                brand_canon["guidelines"].append(sentence.strip())
    
    return brand_canon